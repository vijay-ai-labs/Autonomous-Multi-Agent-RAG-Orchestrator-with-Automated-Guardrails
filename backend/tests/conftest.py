"""Shared fixtures: generated sample files and Phase 3 API/test doubles.

The environment block runs first (before any ``core.config`` import) so the
no-default ``Settings`` fields resolve to harmless test values; nothing here
opens a real connection.
"""

import os

os.environ.setdefault(
    "POSTGRES_URL", "postgresql+asyncpg://user:pass@localhost:5432/rag_test"
)
os.environ.setdefault("QDRANT_URL", "http://localhost:6333")
os.environ.setdefault("REDIS_URL", "redis://localhost:6379/0")
os.environ.setdefault("OPENAI_API_KEY", "sk-test-key")
os.environ.setdefault("JWT_SECRET", "test-jwt-secret")

from pathlib import Path
from uuid import UUID, uuid4

import pytest
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Known marker text used by assertions across the parser tests.
PDF_HEADING = "EMPLOYEE HANDBOOK"
PDF_PAGE1_BODY = "This is the first page of the handbook with introductory content."
PDF_PAGE2_BODY = "This is the second page describing leave and benefits policy."

DOCX_HEADING_1 = "Onboarding Policy"
DOCX_HEADING_2 = "Equipment"
DOCX_TABLE_CELLS = ["Asset", "Owner", "Laptop", "IT"]

HTML_SECTION_A = "Getting Started"
HTML_SECTION_B = "Access Requests"


@pytest.fixture
def base_metadata() -> dict:
    """Metadata dict shared by all parsers (mirrors the Phase 3 contract)."""
    return {
        "filename": "stored-uuid.bin",
        "original_filename": "sample",
        "doc_type": "policy",
        "department": "hr",
        "file_size_bytes": 0,
    }


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    """A real 2-page PDF: page 1 has an ALL-CAPS heading, both pages have body text."""
    path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)

    # Page 1: heading at top, then body.
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, PDF_HEADING)
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, PDF_PAGE1_BODY)
    c.showPage()

    # Page 2: body only.
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, PDF_PAGE2_BODY)
    c.showPage()

    c.save()
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    """A DOCX with two headings, three paragraphs, one empty paragraph, and a 2x2 table."""
    path = tmp_path / "sample.docx"
    doc = Document()

    doc.add_heading(DOCX_HEADING_1, level=1)
    doc.add_paragraph("All new hires must complete onboarding within the first week.")
    doc.add_paragraph("Managers are responsible for scheduling orientation sessions.")
    doc.add_paragraph("")  # empty paragraph — must be skipped
    doc.add_heading(DOCX_HEADING_2, level=2)
    doc.add_paragraph("Company laptops are issued by the IT department.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = DOCX_TABLE_CELLS[0]
    table.cell(0, 1).text = DOCX_TABLE_CELLS[1]
    table.cell(1, 0).text = DOCX_TABLE_CELLS[2]
    table.cell(1, 1).text = DOCX_TABLE_CELLS[3]

    doc.save(str(path))
    return path


@pytest.fixture
def html_file(tmp_path: Path) -> Path:
    """A Confluence-style export with nav/footer/script/style/metadata noise."""
    path = tmp_path / "sample.html"
    html = f"""<!DOCTYPE html>
<html>
<head>
  <style>.x {{ color: red; }}</style>
  <script>var tracking = "should-not-appear";</script>
</head>
<body>
  <nav>SHOULD_NOT_APPEAR_NAV</nav>
  <header>SHOULD_NOT_APPEAR_HEADER</header>
  <div class="breadcrumb">SHOULD_NOT_APPEAR_BREADCRUMB</div>
  <div class="page-metadata">Created by jdoe SHOULD_NOT_APPEAR_META</div>
  <div role="navigation">SHOULD_NOT_APPEAR_ROLE_NAV</div>
  <main>
    <h2>{HTML_SECTION_A}</h2>
    <p>Welcome to the internal knowledge base for engineering.</p>
    <ul>
      <li>Request an account from IT</li>
      <li>Read the security policy</li>
    </ul>
    <h2>{HTML_SECTION_B}</h2>
    <p>Submit access requests through the service portal.</p>
  </main>
  <footer id="footer">SHOULD_NOT_APPEAR_FOOTER</footer>
</body>
</html>"""
    path.write_text(html, encoding="utf-8")
    return path


# --------------------------------------------------------------------------
# Phase 3: auth + DB test doubles
# --------------------------------------------------------------------------


class FakeResult:
    """Mimics the subset of AsyncResult that route handlers use after execute()."""

    def __init__(self, rows: list | None = None) -> None:
        self._rows = rows or []

    def scalar(self):
        return self._rows[0] if self._rows else None

    def scalar_one_or_none(self):
        return self._rows[0] if self._rows else None

    def scalars(self):
        return self

    def all(self) -> list:
        return list(self._rows)

    def first(self):
        return self._rows[0] if self._rows else None

    def __iter__(self):
        return iter(self._rows)


class FakeAsyncSession:
    """Minimal stand-in for AsyncSession used by API route tests.

    Records added objects and assigns a UUID id on ``refresh`` (mirroring the
    Postgres ``gen_random_uuid()`` default) so routes can return ``document.id``.
    """

    def __init__(self) -> None:
        self.added: list = []
        # Tests can pre-seed rows: session._results[query_key] = [row, ...]
        self._results: dict = {}

    def add(self, obj) -> None:
        self.added.append(obj)

    def add_all(self, objs) -> None:
        self.added.extend(objs)

    async def commit(self) -> None:
        return None

    async def flush(self) -> None:
        return None

    async def refresh(self, obj) -> None:
        if getattr(obj, "id", None) is None:
            obj.id = uuid4()

    async def get(self, model, pk):
        return None

    async def execute(self, *args, **kwargs) -> FakeResult:
        return FakeResult()


@pytest.fixture
def user_id() -> UUID:
    """A stable authenticated user id for token generation."""
    return uuid4()


@pytest.fixture
def auth_headers(user_id: UUID) -> dict:
    """Authorization header carrying a valid signed JWT for ``user_id``."""
    from api.middleware.auth import create_access_token

    return {"Authorization": f"Bearer {create_access_token(user_id)}"}


@pytest.fixture
def fake_db() -> FakeAsyncSession:
    """A fresh fake async DB session per test."""
    return FakeAsyncSession()


@pytest.fixture
def patch_scope(monkeypatch, user_id: UUID):
    """Patch RBAC scope resolution so route tests skip DB/Redis lookups.

    Returns a setter so a test can choose the role/departments; defaults to an
    ``admin`` scope (sees everything) to match pre-RBAC route-test expectations.
    The bearer-token check still runs (``get_current_user`` layers on
    ``get_current_user_id``), so missing/invalid-auth tests behave unchanged.
    """
    from api.middleware import auth
    from core.access import UserScope

    state: dict = {"role": "admin", "departments": ()}

    async def _fake_load_scope(uid: UUID) -> UserScope:
        return UserScope(id=uid, role=state["role"], departments=tuple(state["departments"]))

    monkeypatch.setattr(auth, "_load_scope", _fake_load_scope)

    def _set(role: str, departments: tuple[str, ...] = ()) -> None:
        state["role"] = role
        state["departments"] = departments

    return _set
