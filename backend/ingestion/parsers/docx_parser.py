"""DOCX parser built on python-docx (>= 1.1).

DOCX has no real page concept, so paragraphs (and tables) are streamed in document
order and grouped into ~3000-char synthetic pages. The section assigned to a page is
the last ``Heading``-styled paragraph seen before that page begins.
"""

import logging
from pathlib import Path

from docx import Document

from ingestion.parsers.base import BaseParser, clean_whitespace, group_into_pages
from ingestion.schemas import ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)

# Sentinel prepended to a line to mark it as a heading boundary so we can recover
# the section for each synthetic page after grouping. Uses a control char unlikely
# to appear in document text.
_HEADING_MARK = "\x00HEADING\x00"


class DocxParser(BaseParser):
    """Extract paragraphs and tables, grouping them into synthetic pages."""

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path, metadata: dict) -> ParsedDocument:
        document = Document(str(file_path))

        lines: list[str] = []
        empty_skipped = 0

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                empty_skipped += 1
                continue
            style_name = para.style.name if para.style is not None else ""
            if style_name.startswith("Heading"):
                lines.append(f"{_HEADING_MARK}{text}")
            else:
                lines.append(text)

        # Tables: each row's cells joined by " | ", rows joined by newline.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))

        if empty_skipped:
            logger.debug("Skipped %d empty paragraph(s) in %s", empty_skipped, file_path.name)

        stream = "\n".join(lines)
        groups = group_into_pages(stream)

        pages: list[ParsedPage] = []
        last_section: str | None = None
        for index, group in enumerate(groups, start=1):
            # The section for this page is the heading in effect when it starts:
            # the most recent heading from a prior page carries over.
            page_section = last_section
            cleaned_lines: list[str] = []
            for line in group.split("\n"):
                if line.startswith(_HEADING_MARK):
                    heading = line[len(_HEADING_MARK):].strip()
                    last_section = heading
                    if page_section is None:
                        page_section = heading
                    cleaned_lines.append(heading)
                else:
                    cleaned_lines.append(line)
            content = clean_whitespace("\n".join(cleaned_lines))
            pages.append(
                ParsedPage(page_number=index, content=content, section=page_section)
            )

        return ParsedDocument(
            filename=metadata["filename"],
            original_filename=metadata["original_filename"],
            doc_type=metadata["doc_type"],
            department=metadata.get("department"),
            pages=pages,
            page_count=len(pages),
            file_size_bytes=metadata["file_size_bytes"],
        )
